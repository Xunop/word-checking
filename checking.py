from rules import *
from utils import is_cjk_char, is_latin_char, is_punctuation, get_effective_run_fonts

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.document import Document as DocObject  # For type hinting
from docx.styles.style import _ParagraphStyle  # For type hinting
import logging
import re
import sys
import html
from font import get_effective_font_property

logging.basicConfig(
    format="{levelname} - {message}", style="{", level=logging.INFO
)


class FormatChecker:
    def __init__(self, rules):
        self.rules = rules
        self.errors = []
        self.default_style_name = self._find_default_style_name()

    def _find_default_style_name(self):
        for name, style_rules in self.rules["paragraph"].items():
            if style_rules.get("is_default", False):
                return name
        if "论文正文" in self.rules["paragraph"]:
            return "论文正文"
        if "Normal" in self.rules["paragraph"]:
            return "Normal"
        return None

    def _add_error(
        self,
        para_idx,
        style_name,
        paragraph_main_snippet,
        full_paragraph_text,
        error_category,
        rule_key,
        expected,
        actual,
        run_idx=None,
        run_text_snippet_for_detail=None,
        error_char_location=None,
    ):
        para_error_block = None
        for block in self.errors:
            if block["para_idx"] == para_idx:
                para_error_block = block
                break

        if para_error_block is None:
            para_error_block = {
                "para_idx": para_idx,
                "style_name": style_name,
                "paragraph_text_snippet": paragraph_main_snippet,
                "full_text": full_paragraph_text,
                "details": [],
            }
            self.errors.append(para_error_block)
            self.errors.sort(key=lambda x: x["para_idx"])

        error_item = {
            "category": error_category,
            "rule": rule_key,
            "expected": str(expected),
            "actual": str(actual),
        }
        if run_idx is not None:
            error_item["run_idx"] = run_idx
            error_item["run_text"] = (
                run_text_snippet_for_detail if run_text_snippet_for_detail else ""
            )

        if error_char_location:
            error_item["location"] = error_char_location

        para_error_block["details"].append(error_item)

        # log_msg_for_debug = (
        #     f"Para {para_idx+1} (Style: '{style_name}', Snippet: '{paragraph_main_snippet[:20]}...') - "
        #     f"Category: '{error_category}', Rule: '{rule_key}', Expected: {expected}, Actual: {actual}"
        # )
        # if run_idx is not None:
        #     log_msg_for_debug += f", Run {run_idx+1}"
        #     if run_text_snippet_for_detail:
        #         log_msg_for_debug += f" ('{run_text_snippet_for_detail[:15]}...')"
        # if error_char_location:
        #     log_msg_for_debug += f", Location: {error_char_location}"
        # logging.debug(f"Adding structured error: {log_msg_for_debug}")

    def _get_effective_format_value(
        self, direct_format, style_format, attribute_name, default_value
    ):
        """
        获取有效的格式值。
        首先检查直接格式，然后检查样式格式，最后使用默认值。
        """
        value = getattr(direct_format, attribute_name, None)
        if value is None and style_format:
            value = getattr(style_format, attribute_name, None)

        if value is None:
            return default_value
        return value

    def _get_first_line_location(self, paragraph_text):
        """获取段落首行索引"""
        if not paragraph_text:
            return None
        newline_index = paragraph_text.find('\n')
        if newline_index == -1: # No newline, whole text is one line
            return [0, len(paragraph_text)]
        else:
            return [0, newline_index]


    def get_effective_rules(self, style_name_or_obj):
        """
        获取指定样式的有效规则，处理 based_on 继承和全局默认。
        """
        # 1. 从全局字体和间距规则开始
        effective_rules = {}
        effective_rules.update(self.rules.get("fonts", {}))
        effective_rules.update(self.rules.get("spacing", {}))
        effective_rules.update(self.rules.get("section", {}))

        style_name = ""
        if isinstance(style_name_or_obj, str):
            style_name = style_name_or_obj
        elif isinstance(style_name_or_obj, _ParagraphStyle):
            style_name = style_name_or_obj.name

        # 尝试获取特定样式规则，如果找不到，并且是Word的"Normal"（正文）样式，则使用配置中的默认样式
        style_to_check = self.rules["paragraph"].get(style_name)
        if not style_to_check and style_name in ["Normal", "正文"] and self.default_style_name:
            style_to_check = self.rules["paragraph"].get(
                self.default_style_name)
            if style_to_check:
                logging.info(f"Style '{style_name}' not in rules, using default style '{
                             self.default_style_name}' for checking.")
                style_name = self.default_style_name  # 更新style_name为实际使用的规则名

        if not style_to_check:
            logging.warning(f"Style '{
                            style_name}' not found in DEFAULT_RULES['paragraph'] and no default mapping applied.")
            return effective_rules  # 只返回全局规则

        # 2. 处理 based_on 继承链
        style_chain = []
        current_name = style_name
        while current_name:
            current_style_rules = self.rules["paragraph"].get(current_name)
            if not current_style_rules:
                logging.warning(
                    f"Base style '{current_name}' not found in rules.")
                break
            style_chain.insert(0, current_style_rules)  # 从基样式到当前样式
            current_name = current_style_rules.get("based_on")

        # 3. 合并规则，子样式覆盖父样式
        for style_in_chain in style_chain:
            effective_rules.update(style_in_chain)

        return effective_rules

    def check_paragraph_formatting(self, p, p_idx, effective_rules, style_name):
        para_text_snippet = p.text[:30].replace("\n", " ")
        # highlighting context
        full_para_text = p.text
        first_line_loc = self._get_first_line_location(full_para_text)
        direct_fmt = p.paragraph_format
        style_p_fmt = (
            p.style.paragraph_format if p.style else None
        )  # 获取段落样式的格式

        # 调试输出，可以保留或删除
        # print(f"--- 段落 {p_idx} ({para_text_snippet}) ---")
        # print(f"直接格式行距: {direct_fmt.line_spacing}, 规则: {direct_fmt.line_spacing_rule}")
        # if style_p_fmt:
        #     print(f"样式格式行距: {style_p_fmt.line_spacing}, 规则: {style_p_fmt.line_spacing_rule}")
        # else:
        #     print("无样式格式对象")

        if "alignment" in effective_rules:
            actual_alignment = self._get_effective_format_value(
                direct_fmt, style_p_fmt, "alignment", WD_ALIGN_PARAGRAPH.LEFT
            )
            if actual_alignment != effective_rules["alignment"]:
                self._add_error(
                    p_idx,
                    style_name,
                    para_text_snippet,
                    full_para_text,
                    "段落格式",
                    "alignment",
                    WD_ALIGN_PARAGRAPH(effective_rules["alignment"]).name,
                    (
                        WD_ALIGN_PARAGRAPH(actual_alignment).name
                        if actual_alignment is not None
                        else "None"
                    ),
                    error_char_location=first_line_loc
                )

        if "first_line_indent_pt" in effective_rules:
            expected_indent_emu = effective_rules["first_line_indent_pt"]
            actual_indent_raw = self._get_effective_format_value(
                direct_fmt, style_p_fmt, "first_line_indent", 0
            )
            actual_indent_emu = (
                actual_indent_raw.pt if actual_indent_raw is not None and actual_indent_raw != 0 else 0
            )  # 确保是数值

            if abs(actual_indent_emu - expected_indent_emu) > PT_TOLERANCE:
                self._add_error(
                    p_idx,
                    style_name,
                    para_text_snippet,
                    full_para_text,
                    "段落格式",
                    "first_line_indent_pt",
                    f"{effective_rules['first_line_indent_pt']:.2f} pt",
                    f"{actual_indent_emu:.2f} pt",
                    error_char_location=first_line_loc
                )

        # 行间距规则
        actual_ls_rule = None
        if "line_spacing_rule" in effective_rules:
            actual_ls_rule = self._get_effective_format_value(
                direct_fmt,
                style_p_fmt,
                "line_spacing_rule",
                WD_LINE_SPACING.SINGLE,  # 默认单倍行距
            )
            if actual_ls_rule != effective_rules["line_spacing_rule"]:
                self._add_error(
                    p_idx,
                    style_name,
                    para_text_snippet,
                    full_para_text,
                    "段落格式",
                    "line_spacing_rule",
                    WD_LINE_SPACING(effective_rules["line_spacing_rule"]).name,
                    (
                        WD_LINE_SPACING(actual_ls_rule).name
                        if actual_ls_rule is not None
                        else "None"
                    ),
                    error_char_location=first_line_loc
                )

        # 行间距值 (仅当规则匹配或规则允许自定义值时检查)
        # 注意：actual_ls_rule 此处使用的是已经经过回退逻辑确定的规则
        if (
            "line_spacing_value" in effective_rules
            and actual_ls_rule
            == effective_rules.get(
                "line_spacing_rule"
            )  # 确保规则是我们期望检查值的规则
            and effective_rules["line_spacing_rule"]
            in [
                WD_LINE_SPACING.MULTIPLE,
                WD_LINE_SPACING.AT_LEAST,
                WD_LINE_SPACING.EXACTLY,
            ]
        ):

            expected_val = float(effective_rules["line_spacing_value"])
            # 对于行距值，如果规则是 SINGLE, DOUBLE, ONE_POINT_FIVE，其 line_spacing 属性可能为 None 或特定值 (如 1.0, 2.0, 1.5)
            # 当规则是 MULTIPLE, AT_LEAST, EXACTLY 时，line_spacing 存储的是 Pt 值 (对于 MULTIPLE，是浮点数)

            actual_val_raw = self._get_effective_format_value(
                direct_fmt, style_p_fmt, "line_spacing", None  # 先获取原始值
            )

            actual_val = 0.0  # 默认值
            if actual_val_raw is not None:
                if effective_rules["line_spacing_rule"] == WD_LINE_SPACING.MULTIPLE:
                    # python-docx 对于多倍行距，line_spacing 直接返回浮点数 (例如 1.5 代表1.5倍行距)
                    actual_val = (
                        float(actual_val_raw / 127000)
                        if isinstance(actual_val_raw, int) and actual_val_raw > 200
                        else float(actual_val_raw)
                    )
                elif effective_rules["line_spacing_rule"] in [
                    WD_LINE_SPACING.AT_LEAST,
                    WD_LINE_SPACING.EXACTLY,
                ]:
                    # 对于 AT_LEAST 和 EXACTLY，line_spacing 返回的是 EMU，需要转换为 Pt
                    actual_val = (
                        Pt(actual_val_raw).pt
                        if isinstance(actual_val_raw, int)
                        else float(actual_val_raw)
                    )  # 假设如果不是int，已经是Pt
                else:  # 其他情况（SINGLE, DOUBLE, ONE_POINT_FIVE），line_spacing 可能不是我们期望的数值比较对象
                    actual_val = (
                        float(actual_val_raw) if actual_val_raw is not None else 1.0
                    )  # 默认给个值避免比较错误

            if abs(actual_val - expected_val) > FLOAT_TOLERANCE:
                self._add_error(
                    p_idx,
                    style_name,
                    para_text_snippet,
                    full_para_text,
                    "段落格式",
                    "line_spacing_value",
                    f"{expected_val:.2f}",
                    f"{actual_val:.2f}",
                    error_char_location=first_line_loc
                )

        if "space_before_pt" in effective_rules:
            expected_pt_val = effective_rules["space_before_pt"]
            space_before_obj = self._get_effective_format_value(
                direct_fmt, style_p_fmt, "space_before", None
            )
            actual_pt = space_before_obj.pt if space_before_obj is not None else 0

            if abs(actual_pt - expected_pt_val) > PT_TOLERANCE:
                self._add_error(
                    p_idx,
                    style_name,
                    para_text_snippet,
                    full_para_text,
                    "段落格式",
                    "space_before_pt",
                    f"{expected_pt_val:.1f} pt",
                    f"{actual_pt:.1f} pt",
                    error_char_location=first_line_loc
                )

        if "space_after_pt" in effective_rules:
            expected_pt_val = effective_rules["space_after_pt"]
            space_after_obj = self._get_effective_format_value(
                direct_fmt, style_p_fmt, "space_after", None
            )
            actual_pt = space_after_obj.pt if space_after_obj is not None else 0

            if abs(actual_pt - expected_pt_val) > PT_TOLERANCE:
                self._add_error(
                    p_idx,
                    style_name,
                    para_text_snippet,
                    full_para_text,
                    "段落格式",
                    "space_after_pt",
                    f"{expected_pt_val:.1f} pt",
                    f"{actual_pt:.1f} pt",
                    error_char_location=first_line_loc
                )

        if "keep_with_next" in effective_rules:
            actual_kwn = self._get_effective_format_value(
                direct_fmt,
                style_p_fmt,
                "keep_with_next",
                False,  # Word 默认可能是 False
            )
            if actual_kwn != effective_rules["keep_with_next"]:
                self._add_error(
                    p_idx,
                    style_name,
                    para_text_snippet,
                    full_para_text,
                    "段落格式",
                    "keep_with_next",
                    effective_rules["keep_with_next"],
                    actual_kwn,
                    error_char_location=first_line_loc
                )

        if "keep_together" in effective_rules:
            actual_kt = self._get_effective_format_value(
                direct_fmt, style_p_fmt, "keep_together", False  # Word 默认可能是 False
            )
            if actual_kt != effective_rules["keep_together"]:
                self._add_error(
                    p_idx,
                    style_name,
                    para_text_snippet,
                    full_para_text,
                    "段落格式",
                    "keep_together",
                    effective_rules["keep_together"],
                    actual_kt,
                    error_char_location=first_line_loc
                )

        if "widow_control" in effective_rules:
            # Word 的默认 widow_control 通常是 True (如果样式中未指定)
            # python-docx 未显式设置时可能返回 None
            actual_wc = self._get_effective_format_value(
                direct_fmt,
                style_p_fmt,
                "widow_control",
                True,
            )
            
            actual_wc_for_comparison = actual_wc if actual_wc is not None else False

            if actual_wc_for_comparison != effective_rules["widow_control"]:
                self._add_error(
                    p_idx,
                    style_name,
                    para_text_snippet,
                    full_para_text,
                    "段落格式",
                    "widow_control",
                    effective_rules["widow_control"],
                    actual_wc_for_comparison,
                    error_char_location=first_line_loc
                )

    def check_font_rules_for_paragraph(
        self, p, p_idx, effective_rules, style_name, doc
    ):
        paragraph_main_snippet = p.text[:30].replace("\n", " ")
        # paragraph_main_snippet = p.text.replace("\n", " ")
        full_para_text = p.text
        current_char_offset_in_para = 0


        for r_idx, run in enumerate(p.runs):
            run_text = run.text
            run_len = len(run_text)
            run_start_offset_in_para = current_char_offset_in_para
            run_end_offset_in_para = run_start_offset_in_para + run_len
            
            # run 在当前段落的位置
            run_loc_in_para = [run_start_offset_in_para, run_end_offset_in_para]

            if not run_text.strip():
                current_char_offset_in_para += run_len
                continue

            
            # print(f"DEBUG::, {run_text}, {run.font.size} {run.style.font.size} {get_effective_font_property(p, run, "size")}")
            # print(f"DEBUG::, {run_text}, {run.font.italic} {run.style.font.italic} {get_effective_font_property(p, run, "italic")}")
            # print(f"DEBUG::, {run_text}, {run.font.italic} {run.style.font.italic} {get_effective_font_property(p, run, "name")}")
            # continue
            run_text_snippet_detail = run.text[:20].replace("\n", " ")
            font = run.font

            if "font_size_pt" in effective_rules:
                expected_size_val = effective_rules["font_size_pt"]
                # actual_size = (
                #     font.size.pt if font.size is not None else 0
                # )
                # if (
                #     actual_size == 0 and p.style.font.size
                # ):  # 如果 run 的字体大小字段为 0/None，检查段落样式字体大小
                #     actual_size = p.style.font.size.pt
                actual_size = get_effective_font_property(p, run, "size")

                if abs(actual_size - expected_size_val) > PT_TOLERANCE:
                    self._add_error(
                        p_idx,
                        style_name,
                        paragraph_main_snippet,
                        full_para_text,
                        "字体",
                        "font_size_pt",
                        f"{expected_size_val:.1f} pt",
                        f"{actual_size:.1f} pt",
                        run_idx=r_idx,
                        run_text_snippet_for_detail=run_text_snippet_detail,
                        error_char_location=run_loc_in_para
                    )

            if "font_bold" in effective_rules:
                # actual_bold = (
                #     font.bold
                #     if font.bold is not None
                #     else (p.style.font.bold if p.style.font else False)
                # )
                actual_bold = get_effective_font_property(p, run, "bold")
                if actual_bold != effective_rules["font_bold"]:
                    self._add_error(
                        p_idx,
                        style_name,
                        paragraph_main_snippet,
                        full_para_text,
                        "字体",
                        "font_bold",
                        effective_rules["font_bold"],
                        actual_bold,
                        run_idx=r_idx,
                        run_text_snippet_for_detail=run_text_snippet_detail,
                        error_char_location=run_loc_in_para
                    )

            if "font_italic" in effective_rules:
                # actual_italic = (
                #     font.italic
                #     if font.italic is not None
                #     else (p.style.font.italic if p.style.font else False)
                # )  # Check style if None
                actual_italic = get_effective_font_property(p, run, "italic")
                if actual_italic != effective_rules["font_italic"]:
                    self._add_error(
                        p_idx,
                        style_name,
                        paragraph_main_snippet,
                        full_para_text,
                        "字体",
                        "font_italic",
                        effective_rules["font_italic"],
                        actual_italic,
                        run_idx=r_idx,
                        run_text_snippet_for_detail=run_text_snippet_detail,
                        error_char_location=run_loc_in_para
                    )

            run_text = run.text
            is_chinese_dominant = bool(re.search(RE_CHINESE, run_text))
            is_western_dominant = bool(re.search(RE_WESTERN, run_text))
            # is_number_dominant = bool(re.search(RE_NUMBER, run_text)) # RE_NUMBER might be too broad if it includes western numbers

            target_font_key = None
            target_font_value = None
            font_to_check_actual = None

            effective_run_fonts = get_effective_run_fonts(run, p, doc)

            if is_chinese_dominant and "chinese_font" in effective_rules:
                target_font_key = "chinese_font"
                target_font_value = effective_rules["chinese_font"]
                font_to_check_actual = effective_run_fonts.get("eastAsia")
            elif (
                is_western_dominant or (not is_chinese_dominant and run_text.strip())
            ) and "western_font" in effective_rules:
                target_font_key = "western_font"
                target_font_value = effective_rules["western_font"]
                font_to_check_actual = effective_run_fonts.get("ascii")
                if font_to_check_actual is None:
                    font_to_check_actual = effective_run_fonts.get("hAnsi")

            # 既不是中文也不是英文则大概率是标点符号（不保证完整），目前暂时先不考虑
            # if not target_font_key and "common_script_font" in effective_rules:
            #     pass

            if target_font_key and target_font_value:
                if font_to_check_actual != target_font_value:
                    normalized_actual = (
                        font_to_check_actual.replace(" (正文)", "").replace(
                            " (标题)", ""
                        )
                        if font_to_check_actual
                        else None
                    )
                    normalized_expected = (
                        target_font_value.replace(" (正文)", "").replace(" (标题)", "")
                        if target_font_value
                        else None
                    )
                    if normalized_actual != normalized_expected:
                        self._add_error(
                            p_idx,
                            style_name,
                            paragraph_main_snippet,
                            full_para_text,
                            "字体",
                            target_font_key,
                            target_font_value,
                            font_to_check_actual,
                            run_idx=r_idx,
                            run_text_snippet_for_detail=run_text_snippet_detail,
                            error_char_location=run_loc_in_para
                        )
            elif (
                "western_font" in effective_rules
                and not is_chinese_dominant
                and run_text.strip()
            ):  # Default check for western_font if no other dominant type matched.
                # This is a broader fallback check if no specific script was dominant.
                actual_font_name = effective_run_fonts.get(
                    "ascii", effective_run_fonts.get("hAnsi")
                )
                if actual_font_name != effective_rules["western_font"]:
                    self._add_error(
                        p_idx,
                        style_name,
                        paragraph_main_snippet,
                        full_para_text,
                        "字体",
                        "western_font (fallback)",
                        effective_rules["western_font"],
                        actual_font_name,
                        run_idx=r_idx,
                        run_text_snippet_for_detail=run_text_snippet_detail,
                        error_char_location=run_loc_in_para
                    )
            current_char_offset_in_para += run_len

    def check_spacing_rules_for_paragraph(self, p, p_idx, effective_rules, style_name):
        text = p.text
        paragraph_main_snippet = text[:50].replace("\n", " ")
        # paragraph_main_snippet = text.replace("\n", " ")
        error_category = "内容间距"

        rule_cn_en = effective_rules.get("require_space_between_cn_en")
        if rule_cn_en is not None:
            if rule_cn_en is True: # 需要空格, 发现没有空格
                for pat, desc in [(f"({RE_CHINESE})({RE_WESTERN})", "(中->英)"), (f"({RE_WESTERN})({RE_CHINESE})", "(英->中)")]:
                    for match in re.finditer(pat, text):
                        loc = [match.start(0), match.end(0)] # Highlight "好w"
                        self._add_error(p_idx, style_name, paragraph_main_snippet, text, error_category,
                                        f"require_space_between_cn_en {desc}", "需要空格", "无空格", error_char_location=loc)
            else: # 不需要空格，但是发现存在空格
                for pat, desc in [(f"({RE_CHINESE})(\\s+)({RE_WESTERN})", "(中->英)"), (f"({RE_WESTERN})(\\s+)({RE_CHINESE})", "(英->中)")]:
                    for match in re.finditer(pat, text):
                        loc = [match.start(2), match.end(2)] # Highlight the space(s)
                        self._add_error(p_idx, style_name, paragraph_main_snippet, text, error_category,
                                        f"require_space_between_cn_en {desc}", "不允许空格", "有空格", error_char_location=loc)


        # Rule: require_space_between_cn_number
        rule_cn_num = effective_rules.get("require_space_between_cn_number")
        if rule_cn_num is not None:
            if rule_cn_num is True:
                for pat, desc in [(f"({RE_CHINESE})({RE_NUMBER})", "(中->数)"), (f"({RE_NUMBER})({RE_CHINESE})", "(数->中)")]:
                    for match in re.finditer(pat, text):
                        loc = [match.start(0), match.end(0)]
                        self._add_error(p_idx, style_name, paragraph_main_snippet, text, error_category,
                                        f"require_space_between_cn_number {desc}", "需要空格", "无空格", error_char_location=loc)
            else:
                for pat, desc in [(f"({RE_CHINESE})(\\s+)({RE_NUMBER})", "(中->数)"), (f"({RE_NUMBER})(\\s+)({RE_CHINESE})", "(数->中)")]:
                    for match in re.finditer(pat, text):
                        loc = [match.start(2), match.end(2)]
                        self._add_error(p_idx, style_name, paragraph_main_snippet, text, error_category,
                                        f"require_space_between_cn_number {desc}", "不允许空格", "有空格", error_char_location=loc)

        # Rule: require_space_between_en_number (similar logic)
        rule_en_num = effective_rules.get("require_space_between_en_number")
        if rule_en_num is not None:
            if rule_en_num is True:
                for pat, desc in [(f"({RE_WESTERN})({RE_NUMBER})", "(英->数)"), (f"({RE_NUMBER})({RE_WESTERN})", "(数->英)")]:
                    for match in re.finditer(pat, text):
                        loc = [match.start(0), match.end(0)]
                        self._add_error(p_idx, style_name, paragraph_main_snippet, text, error_category,
                                        f"require_space_between_en_number {desc}", "需要空格", "无空格", error_char_location=loc)
            else:
                for pat, desc in [(f"({RE_WESTERN})(\\s+)({RE_NUMBER})", "(英->数)"), (f"({RE_NUMBER})(\\s+)({RE_WESTERN})", "(数->英)")]:
                    for match in re.finditer(pat, text):
                        loc = [match.start(2), match.end(2)]
                        self._add_error(p_idx, style_name, paragraph_main_snippet, text, error_category,
                                        f"require_space_between_en_number {desc}", "不允许空格", "有空格", error_char_location=loc)
        
        # Rule: space_after_chinese_punctuation
        rule_after_cn_punct = effective_rules.get("space_after_chinese_punctuation")
        if rule_after_cn_punct == "none":
            pat = f"({RE_CHINESE_PUNCTUATION})(\\s+)"
            for match in re.finditer(pat, text):
                loc = [match.start(2), match.end(2)] # Highlight the space(s)
                self._add_error(p_idx, style_name, paragraph_main_snippet, text, error_category,
                                "space_after_chinese_punctuation", "none (无空格)", "有空格", error_char_location=loc)

        # Rule: no_space_around_full_width_brackets
        rule_no_space_brackets = effective_rules.get("no_space_around_full_width_brackets")
        if rule_no_space_brackets is True:
            # Left bracket + space(s)
            pat_left = f"({RE_FULL_WIDTH_BRACKETS_LEFT})(\\s+)"
            for match in re.finditer(pat_left, text):
                loc = [match.start(2), match.end(2)]
                self._add_error(p_idx, style_name, paragraph_main_snippet, text, error_category,
                                "no_space_around_full_width_brackets (左括号后)", "括号内侧无空格", "有空格", error_char_location=loc)
            # Space(s) + Right bracket
            pat_right = f"(\\s+)({RE_FULL_WIDTH_BRACKETS_RIGHT})"
            for match in re.finditer(pat_right, text):
                loc = [match.start(1), match.end(1)]
                self._add_error(p_idx, style_name, paragraph_main_snippet, text, error_category,
                                "no_space_around_full_width_brackets (右括号前)", "括号内侧无空格", "有空格", error_char_location=loc)

        # Rule: no_space_after_full_width_punctuation_to_en_num
        rule_no_space_punct_en_num = effective_rules.get("no_space_after_full_width_punctuation_to_en_num")
        if rule_no_space_punct_en_num is True:
            pat = f"({RE_CHINESE_PUNCTUATION})(\\s+)([{RE_WESTERN.strip('[]')}{RE_NUMBER.strip('[]')}])"
            for match in re.finditer(pat, text):
                loc = [match.start(2), match.end(2)] # Highlight the space(s)
                self._add_error(p_idx, style_name, paragraph_main_snippet, text, error_category,
                                "no_space_after_full_width_punctuation_to_en_num", "全角标点后接英文/数字时无空格", "有空格", error_char_location=loc)

    def check_document(self, doc_path):
        self.errors = []
        try:
            doc = Document(doc_path)
        except Exception as e:
            # For this kind of error, we can't use the structured approach as it's a global doc error
            # We can append a special error dict or just a string message
            self.errors.append(
                {
                    "para_idx": -1,  # Special index for document-level errors
                    "style_name": "N/A",
                    "paragraph_text_snippet": f"无法打开或读取文档 '{doc_path}'",
                    "details": [
                        {
                            "category": "文档读取",
                            "rule": "文件访问",
                            "expected": "成功读取",
                            "actual": f"失败: {e}",
                        }
                    ],
                }
            )
            return self.errors  # Return early

        for i, section in enumerate(doc.sections):
            left = section.left_margin.cm
            right = section.right_margin.cm
            top = section.top_margin.cm
            bottom = section.bottom_margin.cm
            expect_margin = self.rules.get("section")
            expect_left = expect_margin.get("left_margin_cm")
            expect_right = expect_margin.get("right_margin_cm")
            expect_top = expect_margin.get("top_margin_cm")
            expect_bottom = expect_margin.get("bottom_margin_cm")
            if abs(left - expect_left) > CM_TOLERANCE:
                self._add_error(
                    0,
                    None,
                    None,
                    None,
                    "节格式",
                    "页面大小",
                    expect_left,
                    left,
                )
            if abs(right - expect_right) > CM_TOLERANCE:
                self._add_error(
                    0,
                    None,
                    None,
                    None,
                    "节格式",
                    "页面大小",
                    expect_right,
                    right,
                )
            if abs(top - expect_top) > CM_TOLERANCE:
                self._add_error(
                    0,
                    None,
                    None,
                    None,
                    "节",
                    "页面大小",
                    expect_top,
                    top,
                )
            if abs(bottom - expect_bottom) > CM_TOLERANCE:
                self._add_error(
                    0,
                    None,
                    None,
                    None,
                    "节格式",
                    "页面大小",
                    expect_bottom,
                    bottom,
                )

        for p_idx, p in enumerate(doc.paragraphs):
            if not p.text.strip() and not p.runs:
                continue

            style_name = p.style.name
            logging.debug(f"样式名称：{style_name}")
            effective_rules = self.get_effective_rules(p.style)

            # 检查这个段落的样式是否在规则集中，如果没有则回退
            is_style_explicitly_defined = style_name in self.rules["paragraph"] or (
                style_name in ["Normal", "正文"]
                and self.default_style_name in self.rules["paragraph"]
            )

            if (
                not is_style_explicitly_defined
                and self.default_style_name != style_name
                and style_name
                not in self.rules["paragraph"]
                .get(self.default_style_name, {})
                .get("aliases", [])
            ):
                if p.text.strip():
                    logging.info(
                        f"提醒: 段落 {p_idx+1} 使用的样式 '{style_name}' 未在 DEFAULT_RULES 中明确定义，也未映射到默认样式。将仅应用全局规则（如有）。"
                    )

            logging.debug(f"规则集：{effective_rules}")
            if effective_rules:
                self.check_paragraph_formatting(p, p_idx, effective_rules, style_name)
                self.check_font_rules_for_paragraph(
                    p, p_idx, effective_rules, style_name, doc
                )
                self.check_spacing_rules_for_paragraph(
                    p, p_idx, effective_rules, style_name
                )
        return self.errors


    def _generate_highlighted_html_snippet(self, full_text, location, context_chars=20):
        if not location or not full_text:
            return ""
        
        start, end = location
        start = max(0, min(start, len(full_text)))
        end = max(0, min(end, len(full_text)))

        if start > end : return "[无效位置]"
        if start == end:
             pass


        text_len = len(full_text)
        snippet_display_start = max(0, start - context_chars)
        snippet_display_end = min(text_len, end + context_chars)

        prefix = html.escape(full_text[snippet_display_start:start])
        highlighted_content = html.escape(full_text[start:end])
        suffix = html.escape(full_text[end:snippet_display_end])
        
        ellipsis_start = "..." if snippet_display_start > 0 else ""
        ellipsis_end = "..." if snippet_display_end < text_len else ""
        
        # 如果高亮部分为空，使用“空内容或仅空白表示”
        if not prefix and not highlighted_content and not suffix and (ellipsis_start or ellipsis_end):
            if not full_text[snippet_display_start:snippet_display_end].strip():
                 return f"{ellipsis_start}[空内容或仅空白]{ellipsis_end}"

        return f"{ellipsis_start}{prefix}<span class='char-highlight'>{highlighted_content}</span>{suffix}{ellipsis_end}"

    def _generate_highlighted_console_snippet(self, full_text, location, context_chars=20, is_tty=True, colors_class=None):
        if not location or not full_text:
            return ""

        start, end = location
        start = max(0, min(start, len(full_text)))
        end = max(0, min(end, len(full_text)))

        if start > end : return "[无效位置]"
        
        text_len = len(full_text)
        snippet_display_start = max(0, start - context_chars)
        snippet_display_end = min(text_len, end + context_chars)

        prefix = full_text[snippet_display_start:start]
        highlighted_content = full_text[start:end]
        suffix = full_text[end:snippet_display_end]

        ellipsis_start = "..." if snippet_display_start > 0 else ""
        ellipsis_end = "..." if snippet_display_end < text_len else ""

        if not prefix and not highlighted_content and not suffix and (ellipsis_start or ellipsis_end):
            if not full_text[snippet_display_start:snippet_display_end].strip():
                 return f"{ellipsis_start}[空内容或仅空白]{ellipsis_end}"

        if is_tty and colors_class:
            highlighted_part_colored = f"{colors_class.HIGHLIGHT_CHAR}{highlighted_content}{colors_class.ENDC}"
        else:
            highlighted_part_colored = highlighted_content
        
        return f"{ellipsis_start}{prefix}{highlighted_part_colored}{suffix}{ellipsis_end}"


    def generate_html_report(self, filename="format_report.html"):
        html_start = """
        <html><head><meta charset='UTF-8'><title>格式检查报告</title>
        <style>
            body { font-family: 'Segoe UI', Arial, sans-serif; margin: 20px; background-color: #f4f4f4; color: #333; }
            h1 { color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }
            .document-error { background-color: #e74c3c; color: white; padding: 15px; margin-bottom: 20px; border-radius: 5px; }
            .paragraph-errors { margin-bottom: 25px; border: 1px solid #bdc3c7; border-radius: 5px; background-color: #fff; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }
            .paragraph-header { font-size: 1.3em; font-weight: bold; margin-bottom: 15px; color: #3498db; padding: 10px 15px; background-color: #ecf0f1; border-bottom: 1px solid #bdc3c7; border-top-left-radius: 5px; border-top-right-radius: 5px;}
            .paragraph-header .style-name { font-weight: normal; color: #7f8c8d; font-size: 0.9em; }
            .paragraph-header .snippet { font-weight: normal; font-style: italic; color: #555; font-size: 0.9em; display: block; margin-top: 5px;}
            table { width: 100%; border-collapse: collapse; margin-top: 0px; }
            th, td { border-bottom: 1px solid #ddd; padding: 12px 15px; text-align: left; font-size: 0.95em; vertical-align: top;}
            th { background-color: #f8f9fa; color: #34495e; font-weight: 600;}
            tr:last-child td { border-bottom: none; }
            /* tr:hover { background-color: #f1f1f1; } */
            .error-category { font-weight: 500; color: #8e44ad; }
            .error-rule { color: #7f8c8d; }
            .expected { color: #27ae60; font-weight: 500; }
            .actual { color: #c0392b; font-weight: 500; }
            .run-info { font-size: 0.85em; color: #95a5a6; }
            .char-highlight { background-color: #f1c40f; color: #c0392b; font-weight: bold; padding: 0.1em 0; border-radius: 0.2em;}
            .context-snippet { font-family: 'Courier New', Courier, monospace; font-size: 0.9em; color: #555; display: block; margin-top: 5px; white-space: pre-wrap; word-break: break-all;}
        </style>
        </head><body><h1>格式检查报告</h1>
        """
        html_end = "</body></html>"
        report_body = ""

        doc_level_error_processed = False
        if not self.errors:
            html_content = html_start + "<p>未发现格式问题。</p>" + html_end
            try:
                with open(filename, "w", encoding="utf-8") as f:
                    f.write(html_content)
                print(f"\nHTML报告已生成: {filename}")
            except IOError as e:
                print(f"错误: 无法写入HTML报告文件 '{filename}'. 详细信息: {e}")
            return

        for para_error_block in self.errors:
            if para_error_block['para_idx'] == -1: # Document level error
                doc_level_error_processed = True
                continue

            para_idx = para_error_block['para_idx']
            style_name = para_error_block['style_name']
            snippet = para_error_block['paragraph_text_snippet']
            full_para_text = para_error_block['full_text']

            report_body += f"<div class='paragraph-errors'>\n"
            report_body += f"  <div class='paragraph-header'>段落 {para_idx + 1} <span class='style-name'>(样式: '{style_name}')</span><span class='snippet'>内容预览: '{snippet}...'</span></div>\n"
            report_body += "  <table>\n"
            report_body += "    <tr><th>类别</th><th>规则</th><th>期望值</th><th>实际值</th><th>Run/备注</th><th>上下文/高亮</th></tr>\n"

            for err in para_error_block['details']:
                category = err['category']
                rule = err['rule']
                expected = err['expected']
                actual = err['actual']
                run_info_html = ""
                if "run_idx" in err:
                    run_info_html = f"Run {err['run_idx'] + 1}"
                    if err.get("run_text"):
                        run_info_html += f" ('{html.escape(err['run_text'])}')"
                
                highlighted_snippet_html = ""
                if err.get("location"):
                    highlighted_snippet_html = self._generate_highlighted_html_snippet(full_para_text, err["location"])
                    highlighted_snippet_html = f"<span class='context-snippet'>{highlighted_snippet_html}</span>"


                report_body += f"    <tr>\n"
                report_body += f"      <td><span class='error-category'>{category}</span></td>\n"
                report_body += f"      <td><span class='error-rule'>{rule}</span></td>\n"
                report_body += f"      <td><span class='expected'>{expected}</span></td>\n"
                report_body += f"      <td><span class='actual'>{actual}</span></td>\n"
                report_body += f"      <td><span class='run-info'>{run_info_html}</span></td>\n"
                report_body += f"      <td>{highlighted_snippet_html}</td>\n"
                report_body += f"    </tr>\n"
            report_body += "  </table>\n"
            report_body += "</div>\n"
        
        if not self.errors and not doc_level_error_processed:
             report_body = "<p>未发现格式问题。</p>"

        full_html = html_start + report_body + html_end
        try:
            with open(filename, "w", encoding="utf-8") as f:
                f.write(full_html)
            if not doc_level_error_processed or self.errors :
                print(f"\nHTML报告已生成: {filename}")
        except IOError as e:
            print(f"错误: 无法写入HTML报告文件 '{filename}'. 详细信息: {e}")

    def print_structured_errors_to_console(self):
        class Colors:
            HEADER = '\033[95m'; BLUE = '\033[94m'; GREEN = '\033[92m'
            WARNING = '\033[93m'; FAIL = '\033[91m'; ENDC = '\033[0m'
            BOLD = '\033[1m'; UNDERLINE = '\033[4m'; GREY = '\033[90m'
            HIGHLIGHT_CHAR = '\033[1;31;43m' # Bold, Red text, Yellow background

        is_tty = hasattr(sys.stdout, 'isatty') and sys.stdout.isatty()
        def colorize(text, color_code):
            return f"{color_code}{text}{Colors.ENDC}" if is_tty else text
        
        if not self.errors:
            print("\n--- 控制台输出：未发现格式问题 (基于当前规则) ---")
            return

        print("\n--- 文档格式检查发现以下问题 (控制台详细输出) ---")
        for para_error_block in self.errors:
            if para_error_block['para_idx'] == -1:
                continue

            para_idx = para_error_block['para_idx']
            style_name = para_error_block['style_name']
            snippet = para_error_block['paragraph_text_snippet']
            full_para_text = para_error_block['full_text']


            print(f"\n{colorize(f'▼ 段落 {para_idx + 1}', Colors.BOLD + Colors.HEADER)}")
            print(f"  {colorize('样式:', Colors.BLUE)} '{style_name}'")
            print(f"  {colorize('内容片段:', Colors.BLUE)} '{snippet}...'")
            print(f"  {colorize('发现的错误:', Colors.BLUE)}")

            for i, err in enumerate(para_error_block['details']):
                category = err['category']
                rule = err['rule']
                expected = err['expected']
                actual = err['actual']

                print(f"    {i+1}. {colorize(f'[{category}]', Colors.WARNING + Colors.BOLD)}")
                print(f"       {colorize('规则:', Colors.GREY)} {rule}")
                if "run_idx" in err:
                    run_info = f"Run {err['run_idx'] + 1}"
                    if err.get("run_text"):
                        run_info += f" ('{err['run_text']}')"
                    print(f"       {colorize('位置:', Colors.GREY)} {run_info}")
                print(f"       {colorize('期望:', Colors.GREEN)} {expected}")
                print(f"       {colorize('实际:', Colors.FAIL)} {actual}")
                
                if err.get("location"):
                    highlighted_snippet_console = self._generate_highlighted_console_snippet(
                        full_para_text, err["location"], is_tty=is_tty, colors_class=Colors
                    )
                    print(f"         {colorize('上下文:', Colors.GREY)} {highlighted_snippet_console}")


if __name__ == "__main__":
    from rules import DEFAULT_RULES, CM_TOLERANCE, FLOAT_TOLERANCE, PT_TOLERANCE
    from rules import RE_CHINESE, RE_WESTERN, RE_NUMBER, RE_CHINESE_PUNCTUATION
    from rules import RE_FULL_WIDTH_BRACKETS_LEFT, RE_FULL_WIDTH_BRACKETS_RIGHT

    doc_file_path = (
        "test.docx"
    )

    try:
        doc = Document(doc_file_path)
    except Exception:
        print(f"Test file '{doc_file_path}' not found or invalid. Creating a dummy one.")
        doc = Document()
        # 测试段落格式错误（对齐）
        p1 = doc.add_paragraph("This paragraph should be left aligned.", style='Normal')
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p2 = doc.add_paragraph("Text with a ", style='Normal')
        run_correct_font = p2.add_run("correctly_sized")
        run_correct_font.font.size = Pt(11)
        p2.add_run(" segment, then an ")
        run_wrong_size = p2.add_run("ERROR_SIZE")
        run_wrong_size.font.size = Pt(11)
        p2.add_run(" run, and a ")
        run_wrong_bold = p2.add_run("WRONGLY_BOLD")
        run_wrong_bold.font.bold = True
        p2.add_run(" run.")
        
        doc.add_paragraph("First line of text for para rule.\nSecond line of text.\nThird line.", style='Normal')
        
        doc.save(doc_file_path)

    checker = FormatChecker(DEFAULT_RULES)
    checker.check_document(doc_file_path)

    if checker.errors:
        checker.print_structured_errors_to_console()
        checker.generate_html_report("format_checker_report.html")
    else:
        print(f"\n--- 文档 '{doc_file_path}' 未发现格式问题 (基于当前规则) ---")
        checker.generate_html_report(
            "format_checker_report.html"
        )
